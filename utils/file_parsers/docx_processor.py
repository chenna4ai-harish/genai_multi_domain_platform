"""
utils/file_parsers/docx_processor.py

This module provides utilities for extracting text from DOCX (Word) files.

What Does This Do?
------------------
Extracts plain text from Microsoft Word documents (.docx format).
Preserves paragraph structure and handles tables.

Library Used:
-------------
- python-docx: Pure Python library for reading/writing DOCX files

Installation:
-------------
pip install python-docx

Example Usage:
--------------
from utils.file_parsers.docx_processor import DOCXProcessor

processor = DOCXProcessor()
text = processor.extract_text("document.docx")
print(f"Extracted {len(text)} characters")
"""

from typing import List, Dict, Any
from pathlib import Path
import logging

# Try to import python-docx
try:
    import docx

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Configure logging
logger = logging.getLogger(__name__)


class DOCXProcessor:
    """
    DOCX (Word document) text extraction processor.

    Extracts text from .docx files while preserving structure.
    Handles paragraphs, tables, headers, and footers.

    Features:
    ---------
    - Extract all text from document
    - Extract paragraphs separately
    - Extract tables
    - Get document metadata

    Example:
    --------
    processor = DOCXProcessor()

    # Extract all text
    text = processor.extract_text("report.docx")

    # Extract paragraphs
    paragraphs = processor.extract_paragraphs("report.docx")
    for i, para in enumerate(paragraphs, 1):
        print(f"Paragraph {i}: {para[:50]}...")
    """

    def __init__(self):
        """Initialize DOCX processor."""
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is not installed. Install with: pip install python-docx"
            )
        logger.info("DOCXProcessor initialized")

    def extract_text(self, file_path: str, include_tables: bool = True) -> str:

        """
        Extract all text from a DOCX file.

        Parameters:
        -----------
        file_path : str
            Path to DOCX file
        include_tables : bool, optional
            Whether to include table content (default: True)

        Returns:
        --------
        str:
            Extracted text from document

        Example:
        --------
        text = processor.extract_text("document.docx")
        print(f"Length: {len(text)} characters")
        """

        """Extract text from DOCX with PyMuPDF fallback."""
        try:
            # Try PyMuPDF first (better quality)
            try:
                import fitz
                text = ""
                doc = fitz.open(file_path)
                for page in doc:
                    text += page.get_text() + "\n\n"
                doc.close()

                if text.strip():
                    logger.info("Used PyMuPDF for DOCX extraction")
                    return text.strip()
            except ImportError:
                pass  # Fall back to python-docx

            # Fallback: python-docx
            doc = docx.Document(file_path)

            # Better text extraction (not XML)
            paragraphs_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs_text.append(para.text.strip())

            text = "\n\n".join(paragraphs_text)

            # Extract text from tables if requested
            if include_tables:
                for table in doc.tables:
                    table_text = self._extract_table_text(table)
                    if table_text:
                        text += "\n\n" + table_text

            return text.strip()

        except Exception as e:
            logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
            raise RuntimeError(f"DOCX extraction failed: {e}")

    def extract_paragraphs(self, file_path: str) -> List[str]:
        """
        Extract paragraphs as separate strings.

        Parameters:
        -----------
        file_path : str
            Path to DOCX file

        Returns:
        --------
        List[str]:
            List of paragraph texts

        Example:
        --------
        paragraphs = processor.extract_paragraphs("doc.docx")
        for i, para in enumerate(paragraphs, 1):
            print(f"{i}. {para[:100]}...")
        """
        try:
            doc = docx.Document(file_path)
            return [para.text.strip() for para in doc.paragraphs if para.text.strip()]

        except Exception as e:
            logger.error(f"Failed to extract paragraphs from DOCX {file_path}: {e}")
            raise RuntimeError(f"DOCX paragraph extraction failed: {e}")

    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from DOCX file.

        Returns:
        --------
        dict:
            Metadata including:
            - paragraph_count: Number of paragraphs
            - table_count: Number of tables
            - title: Document title (if available)
            - author: Document author (if available)

        Example:
        --------
        metadata = processor.get_metadata("doc.docx")
        print(f"Paragraphs: {metadata['paragraph_count']}")
        print(f"Tables: {metadata['table_count']}")
        """
        try:
            doc = docx.Document(file_path)

            metadata = {
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'title': doc.core_properties.title if hasattr(doc, 'core_properties') else None,
                'author': doc.core_properties.author if hasattr(doc, 'core_properties') else None
            }

            return metadata

        except Exception as e:
            logger.error(f"Failed to get DOCX metadata {file_path}: {e}")
            return {'paragraph_count': 0, 'table_count': 0, 'title': None, 'author': None}

    def _extract_table_text(self, table) -> str:
        """
        Extract text from a table.

        Parameters:
        -----------
        table : docx.table.Table
            Table object from python-docx

        Returns:
        --------
        str:
            Table content as formatted text
        """
        table_text = []

        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells])
            if row_text.strip():
                table_text.append(row_text)

        return "\n".join(table_text)


# =============================================================================
# USAGE EXAMPLES
# =============================================================================

if __name__ == "__main__":
    """
    Demonstration of DOCXProcessor usage.
    Run: python utils/file_parsers/docx_processor.py
    """

    import logging

    logging.basicConfig(level=logging.INFO, format='%(levelname)s - %(message)s')

    print("=" * 70)
    print("DOCXProcessor Usage Examples")
    print("=" * 70)

    if not DOCX_AVAILABLE:
        print("\n❌ python-docx not installed!")
        print("Install with: pip install python-docx")
        exit(1)

    print("\n1. DOCXProcessor Status")
    print("-" * 70)
    print("✅ python-docx is installed and ready")

    print("\n2. Usage Pattern")
    print("-" * 70)
    print("""
# Initialize processor
processor = DOCXProcessor()

# Extract all text
text = processor.extract_text("document.docx")
print(f"Extracted {len(text)} characters")

# Extract paragraphs separately
paragraphs = processor.extract_paragraphs("document.docx")
for i, para in enumerate(paragraphs, 1):
    print(f"Paragraph {i}: {para[:100]}...")

# Get metadata
metadata = processor.get_metadata("document.docx")
print(f"Paragraphs: {metadata['paragraph_count']}")
print(f"Tables: {metadata['table_count']}")
    """)

    print("\n" + "=" * 70)
    print("DOCXProcessor examples completed!")
    print("=" * 70)
