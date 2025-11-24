"""

utils/file_parsers/docx_processor.py

Phase 2 Enhanced DOCX text extraction with structured output and metadata tracking.

What is This File?
-------------------
DOCX processor that extracts text from Microsoft Word documents with:
- Paragraph-level tracking with character positions
- Table extraction and formatting
- Document metadata (title, author)
- File hash computation for provenance
- Dual extraction strategy (PyMuPDF + python-docx)
- Standardized Phase 2 output format

Why Enhanced for Phase 2?
--------------------------
Phase 2 requires comprehensive metadata for:
- Citations with section numbers: [report.docx:Section 3]
- Character range tracking for precise source location
- Provenance tracking via file hash
- Per-paragraph data with position information

Installation:
-------------
pip install python-docx
pip install PyMuPDF  # Optional but recommended (better quality)

Example Usage (Phase 2):
------------------------
processor = DOCXProcessor()

# Extract with full Phase 2 metadata
result = processor.extract("report.docx")

if result['success']:
    print(f"Text: {result['text'][:100]}...")
    print(f"Paragraphs: {result['metadata']['paragraph_count']}")
    print(f"Method: {result['metadata']['extraction_method']}")
    print(f"Hash: {result['metadata']['file_hash']}")
else:
    print(f"Errors: {result['errors']}")

References:
-----------
- Phase 2 Spec: Section 10.3 (Enhanced File Processors)
- python-docx: https://python-docx.readthedocs.io/
- PyMuPDF: https://pymupdf.readthedocs.io/

"""

from typing import List, Dict, Any, Optional
from pathlib import Path
import logging
import hashlib

# Try importing optional backends
try:
    import fitz  # PyMuPDF

    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

# Import required library
try:
    import docx

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Configure logging
logger = logging.getLogger(__name__)


# =============================================================================
# PHASE 2 FILE HASHING UTILITY
# =============================================================================

def compute_file_hash(file_path: str) -> str:
    """
    Compute SHA-256 hash of file for provenance tracking (Phase 2).

    Parameters:
    -----------
    file_path : str
        Path to file

    Returns:
    --------
    str:
        SHA-256 hash (64 hex characters)
    """
    sha256_hash = hashlib.sha256()

    with open(file_path, 'rb') as f:
        for chunk in iter(lambda: f.read(4096), b''):
            sha256_hash.update(chunk)

    return sha256_hash.hexdigest()


# =============================================================================
# PHASE 2 DOCX PROCESSOR - MAIN CLASS
# =============================================================================

class DOCXProcessor:
    """
    Phase 2 enhanced DOCX processor with standardized output format.

    Extracts text from Microsoft Word documents (.docx) with:
    - Paragraph-level tracking and character positions
    - Table extraction and formatting
    - Document metadata
    - File hash computation
    - Optional PyMuPDF fallback for better quality

    Returns structured Phase 2 format with metadata, section data, and file hash.

    Features:
    ---------
    - Dual extraction (PyMuPDF + python-docx)
    - Graceful fallback on failures
    - Per-paragraph metadata with char ranges
    - Table extraction and formatting
    - File hash for integrity checking
    - Complete error collection

    Example:
    --------
    # Initialize processor
    processor = DOCXProcessor()

    # Extract with Phase 2 format
    result = processor.extract("document.docx")

    # Check success
    if result['success']:
        print(f"Extracted: {len(result['text'])} characters")
        print(f"Paragraphs: {len(result['pages'])}")
        print(f"Hash: {result['metadata']['file_hash']}")
    else:
        print(f"Errors: {result['errors']}")
    """

    def __init__(self):
        """
        Initialize DOCX processor.

        Raises:
        -------
        ImportError:
            If python-docx not installed
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is not installed.\n"
                "Install with: pip install python-docx"
            )

        logger.info(
            f"DOCXProcessor initialized\n"
            f"  python-docx: Available\n"
            f"  PyMuPDF: {'Available' if PYMUPDF_AVAILABLE else 'Not available (optional)'}"
        )

    def extract(
            self,
            file_path: str,
            include_tables: bool = True,
            compute_hash: bool = True
    ) -> Dict[str, Any]:
        """
        Extract text and metadata from DOCX file (Phase 2 standardized format).

        This is the PRIMARY extraction method for Phase 2 service layer.
        Returns structured data with comprehensive metadata and error handling.

        Parameters:
        -----------
        file_path : str
            Path to DOCX file
        include_tables : bool
            Whether to extract table content (default: True)
        compute_hash : bool
            Whether to compute file hash (default: True)

        Returns:
        --------
        Dict[str, Any]:
            {
                'text': str,                    # Full extracted text
                'metadata': {
                    'paragraph_count': int,
                    'table_count': int,
                    'file_hash': str,           # SHA-256 hash
                    'file_size_bytes': int,
                    'extraction_method': str,   # 'pymupdf' or 'python-docx'
                    'title': Optional[str],
                    'author': Optional[str]
                },
                'pages': List[Dict],            # Per-paragraph info (called 'pages' for compatibility)
                'success': bool,                # Extraction successful
                'errors': List[str]             # Any warnings/errors
            }

        Page/Paragraph structure:
        -------------------------
        'pages': [
            {
                'page_num': int,                # Paragraph number (1-indexed)
                'text': str,                    # Paragraph/section text
                'char_range': (start, end),     # Character positions
                'length': int,                  # Text length
                'type': str                     # 'paragraph' or 'table'
            },
            ...
        ]

        Raises:
        -------
        FileNotFoundError:
            If DOCX file doesn't exist
        RuntimeError:
            If extraction completely fails

        Example:
        --------
        processor = DOCXProcessor()
        result = processor.extract("document.docx")

        if result['success']:
            print(f"Paragraphs: {result['metadata']['paragraph_count']}")
            print(f"Method: {result['metadata']['extraction_method']}")

            for para in result['pages']:
                print(f"Para {para['page_num']}: {len(para['text'])} chars")
        else:
            print(f"Errors: {result['errors']}")
        """
        logger.info(f"Extracting DOCX: {file_path}")

        errors = []
        text = ""
        paragraphs_data = []
        extraction_method = "unknown"
        doc_metadata = {}

        # Validate file exists
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        # Get file size
        try:
            file_size = file_path.stat().st_size
        except Exception as e:
            logger.warning(f"Could not get file size: {e}")
            file_size = 0

        # Strategy 1: Try PyMuPDF first (better quality)
        if PYMUPDF_AVAILABLE:
            try:
                logger.debug("Attempting extraction with PyMuPDF...")
                text_pymupdf = self._extract_with_pymupdf(str(file_path))

                if text_pymupdf and text_pymupdf.strip():
                    text = text_pymupdf
                    extraction_method = "pymupdf"
                    # For PyMuPDF, we don't have paragraph-level data
                    # Create single entry
                    paragraphs_data = [{
                        'page_num': 1,
                        'text': text,
                        'char_range': (0, len(text)),
                        'length': len(text),
                        'type': 'document'
                    }]
                    logger.info("✅ Successfully extracted with PyMuPDF")
                else:
                    logger.debug("PyMuPDF returned empty text, trying python-docx...")

            except Exception as e:
                logger.debug(f"PyMuPDF extraction failed: {e}, trying python-docx...")
                errors.append(f"PyMuPDF extraction failed: {str(e)}")

        # Strategy 2: Fallback to python-docx (or primary if PyMuPDF not available)
        if not text:
            try:
                logger.debug("Extracting with python-docx...")
                result_docx = self._extract_with_python_docx(str(file_path), include_tables)
                text = result_docx['text']
                paragraphs_data = result_docx['paragraphs']
                doc_metadata = result_docx['metadata']
                extraction_method = "python-docx"
                logger.info("✅ Successfully extracted with python-docx")

            except Exception as e:
                logger.error(f"python-docx extraction failed: {e}")
                errors.append(f"python-docx extraction failed: {str(e)}")
                raise RuntimeError(
                    f"Failed to extract text from DOCX: {file_path}\n"
                    f"Errors: {'; '.join(errors)}"
                )

        # Compute file hash (Phase 2 requirement)
        file_hash = None
        if compute_hash:
            try:
                file_hash = compute_file_hash(str(file_path))
                logger.debug(f"Computed file hash: {file_hash[:16]}...")
            except Exception as e:
                logger.warning(f"Failed to compute file hash: {e}")
                errors.append(f"Hash computation failed: {str(e)}")
                file_hash = "HASH_FAILED"

        # Build Phase 2 standardized result
        result = {
            'text': text.strip(),
            'metadata': {
                'paragraph_count': len(paragraphs_data),
                'table_count': doc_metadata.get('table_count', 0),
                'file_hash': file_hash,
                'file_size_bytes': file_size,
                'extraction_method': extraction_method,
                'title': doc_metadata.get('title'),
                'author': doc_metadata.get('author')
            },
            'pages': paragraphs_data,  # Called 'pages' for API compatibility
            'success': bool(text) and len(paragraphs_data) > 0,
            'errors': errors
        }

        logger.info(
            f"✅ DOCX extraction complete:\n"
            f"   Method: {extraction_method}\n"
            f"   Paragraphs: {len(paragraphs_data)}\n"
            f"   Text length: {len(text):,} chars\n"
            f"   Errors: {len(errors)}"
        )

        return result

    def extract_text(self, file_path: str, include_tables: bool = True) -> str:
        """
        Extract plain text from DOCX (backward compatibility).

        For Phase 2, use extract() method instead which returns structured data.

        Parameters:
        -----------
        file_path : str
            Path to DOCX file
        include_tables : bool
            Whether to include table content

        Returns:
        --------
        str:
            Extracted text
        """
        result = self.extract(file_path, include_tables, compute_hash=False)
        return result['text']

    def extract_paragraphs(self, file_path: str) -> List[str]:
        """
        Extract paragraphs as separate strings.

        Parameters:
        -----------
        file_path : str
            Path to DOCX file

        Returns:
        --------
        List[str]:
            List of paragraph texts

        Example:
        --------
        paragraphs = processor.extract_paragraphs("doc.docx")
        for i, para in enumerate(paragraphs, 1):
            print(f"{i}. {para[:100]}...")
        """
        try:
            doc = docx.Document(file_path)
            return [para.text.strip() for para in doc.paragraphs if para.text.strip()]

        except Exception as e:
            logger.error(f"Failed to extract paragraphs from DOCX {file_path}: {e}")
            raise RuntimeError(f"DOCX paragraph extraction failed: {e}")

    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from DOCX file.

        Returns:
        --------
        dict:
            Metadata including paragraph_count, table_count, title, author
        """
        try:
            result = self.extract(file_path, compute_hash=False)
            return result['metadata']
        except Exception as e:
            logger.error(f"Metadata extraction failed: {e}")
            return {
                'paragraph_count': 0,
                'table_count': 0,
                'title': None,
                'author': None
            }

    # =========================================================================
    # BACKEND IMPLEMENTATIONS
    # =========================================================================

    def _extract_with_pymupdf(self, file_path: str) -> str:
        """Extract using PyMuPDF (best quality) - BACKEND."""
        import fitz

        text_parts = []
        doc = fitz.open(file_path)

        try:
            for page in doc:
                page_text = page.get_text()
                if page_text.strip():
                    text_parts.append(page_text.strip())
        finally:
            doc.close()

        return "\n\n".join(text_parts)

    def _extract_with_python_docx(
            self,
            file_path: str,
            include_tables: bool = True
    ) -> Dict[str, Any]:
        """
        Extract using python-docx with paragraph tracking.

        Returns structured data with per-paragraph information.
        """
        doc = docx.Document(file_path)

        # Extract paragraphs with position tracking
        paragraphs_data = []
        text_parts = []
        char_position = 0

        # Extract from paragraphs
        for para_idx, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()

            if para_text:
                # Track character range for this paragraph
                char_start = char_position
                char_end = char_position + len(para_text)

                paragraphs_data.append({
                    'page_num': para_idx + 1,
                    'text': para_text,
                    'char_range': (char_start, char_end),
                    'length': len(para_text),
                    'type': 'paragraph'
                })

                text_parts.append(para_text)
                char_position = char_end + 2  # +2 for "\n\n" separator

        # Extract tables if requested
        table_count = 0
        if include_tables:
            for table_idx, table in enumerate(doc.tables):
                table_text = self._extract_table_text(table)

                if table_text:
                    # Track table in paragraphs data
                    char_start = char_position
                    char_end = char_position + len(table_text)

                    paragraphs_data.append({
                        'page_num': len(paragraphs_data) + 1,
                        'text': table_text,
                        'char_range': (char_start, char_end),
                        'length': len(table_text),
                        'type': 'table'
                    })

                    text_parts.append(table_text)
                    char_position = char_end + 2
                    table_count += 1

        # Get document metadata
        title = None
        author = None
        if hasattr(doc, 'core_properties'):
            try:
                title = doc.core_properties.title
                author = doc.core_properties.author
            except:
                pass

        # Join all text
        text = "\n\n".join(text_parts)

        return {
            'text': text,
            'paragraphs': paragraphs_data,
            'metadata': {
                'title': title,
                'author': author,
                'table_count': table_count
            }
        }

    def _extract_table_text(self, table) -> str:
        """
        Extract text from a DOCX table.

        Parameters:
        -----------
        table : docx.table.Table
            Table object from python-docx

        Returns:
        --------
        str:
            Table content as formatted text (rows separated by newlines,
            cells separated by |)
        """
        table_text = []

        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells])
            if row_text.strip():
                table_text.append(row_text)

        return "\n".join(table_text)


# =============================================================================
# USAGE EXAMPLES
# =============================================================================

if __name__ == "__main__":
    """
    Demonstration of Phase 2 DOCXProcessor.
    Run: python utils/file_parsers/docx_processor.py
    """
    import logging

    logging.basicConfig(level=logging.INFO, format='%(levelname)s - %(message)s')

    print("=" * 70)
    print("Phase 2 DOCXProcessor - Enhanced DOCX Text Extraction")
    print("=" * 70)

    if not DOCX_AVAILABLE:
        print("\n❌ python-docx not installed!")
        print("Install with: pip install python-docx")
        exit(1)

    # Check available backends
    print("\n1. Available Backends")
    print("-" * 70)
    print(f"python-docx:  {'✅ Available' if DOCX_AVAILABLE else '❌ Not installed'}")
    print(f"PyMuPDF:      {'✅ Available' if PYMUPDF_AVAILABLE else '❌ Not installed (optional)'}")

    # Usage examples
    print("\n2. Phase 2 Usage Pattern")
    print("-" * 70)
    print("""
# Initialize processor
processor = DOCXProcessor()

# Extract with Phase 2 format
result = processor.extract("document.docx")

if result['success']:
    print(f"Text: {len(result['text'])} characters")
    print(f"Paragraphs: {result['metadata']['paragraph_count']}")
    print(f"Method: {result['metadata']['extraction_method']}")
    print(f"Hash: {result['metadata']['file_hash']}")

    # Access per-paragraph data
    for para in result['pages'][:3]:
        print(f"Para {para['page_num']}: {len(para['text'])} chars")
        print(f"  Position: {para['char_range']}")
        print(f"  Type: {para['type']}")
else:
    print(f"Errors: {result['errors']}")

# Extract plain text (backward compatible)
text = processor.extract_text("document.docx")

# Extract paragraphs separately
paragraphs = processor.extract_paragraphs("document.docx")

# Get metadata
metadata = processor.get_metadata("document.docx")
    """)

    print("\n" + "=" * 70)
    print("✅ Phase 2 DOCXProcessor ready to use!")
    print("=" * 70)
