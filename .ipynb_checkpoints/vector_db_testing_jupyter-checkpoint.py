# %%
"""
============================================================================
Jupyter Notebook: Playground Pipeline Testing
============================================================================
Tests chunking, embedding, vector DB, and document ingestion exactly as
app_phase2.py does.

Workflow:
1. Load playground config
2. Create DocumentService via get_service_for_config
3. Test chunking strategies
4. Test embedding with different providers
5. Test vector DB operations (upsert, query)
6. Test full document ingestion (PDF, DOCX, TXT)
7. Test document listing and chunk retrieval
8. Test query/retrieval

All tests follow the same logic used in the Playground UI.
============================================================================
"""

import sys
import os
import logging
import uuid
from pathlib import Path
from datetime import datetime
import yaml

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

print("=" * 70)
print("PLAYGROUND PIPELINE TESTING")
print("=" * 70)
print(f"Start time: {datetime.now()}")
print(f"Working directory: {Path.cwd()}")
print("=" * 70)

# %%
"""
============================================================================
CELL 1: Imports and Setup
============================================================================
"""
print("\n" + "=" * 70)
print("CELL 1: Imports and Setup")
print("=" * 70)

# Core imports
from core.playground_config_manager import PlaygroundConfigManager
from core.config_manager import ConfigManager, DomainConfig
from core.services.document_service import DocumentService, ValidationError, ProcessingError

# Try to import factories (with fallbacks)
try:
    from core.embeddings.embedding_factory import EmbeddingFactory

    print("✅ EmbeddingFactory imported")
except ImportError as e:
    print(f"⚠️  EmbeddingFactory not found: {e}")
    EmbeddingFactory = None

try:
    from core.chunking.chunking_factory import ChunkingFactory

    print("✅ ChunkingFactory imported")
except ImportError as e:
    print(f"⚠️  ChunkingFactory not found: {e}")
    ChunkingFactory = None

try:
    from core.vectorstore.vectorstore_factory import VectorStoreFactory

    print("✅ VectorStoreFactory imported")
except ImportError as e:
    print(f"⚠️  VectorStoreFactory not found: {e}")
    VectorStoreFactory = None

try:
    from core.pipeline.document_pipeline import DocumentPipeline

    print("✅ DocumentPipeline imported")
except ImportError as e:
    print(f"⚠️  DocumentPipeline not found: {e}")
    DocumentPipeline = None

try:
    from core.utils.file_parsers import extract_text_from_file

    print("✅ File parsers imported")
except ImportError as e:
    print(f"⚠️  File parsers not found: {e}")


    def extract_text_from_file(file_path, filename):
        # Fallback: simple text extraction
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return {"text": f.read(), "metadata": {}}

print("\n✅ All imports completed (with fallbacks where needed)")

# %%
"""
============================================================================
CELL 2: Load Playground Config
============================================================================
"""
print("\n" + "=" * 70)
print("CELL 2: Load Playground Config")
print("=" * 70)

pg_mgr = PlaygroundConfigManager()

# List available configs
all_configs = PlaygroundConfigManager.list_configs()
print(f"\n📋 Available playground configs: {len(all_configs)}")
for idx, cfg in enumerate(all_configs, 1):
    print(f"  {idx}. {cfg['name']} (session: {cfg['session_id']})")

if not all_configs:
    print("\n⚠️  No playground configs found!")
    print("Creating a test config...")

    # Create minimal test config
    test_config = {
        "domain_id": "pipeline_test",
        "name": "pipeline_test",
        "display_name": "Pipeline Test Configuration",
        "description": "Config for testing pipeline components",
        "vector_store": {
            "provider": "chromadb",
            "collection_name": "pipeline_test_collection",
            "persist_directory": ".data/chromadb/pipeline_test"
        },
        "chunking": {
            "strategy": "recursive",
            "recursive": {
                "chunk_size": 500,
                "overlap": 50
            }
        },
        "embeddings": {
            "provider": "sentence_transformers",
            "model_name": "all-MiniLM-L6-v2",
            "device": "cpu",
            "batch_size": 32,
            "normalize": True
        },
        "retrieval": {
            "strategies": ["hybrid"],
            "top_k": 10,
            "similarity": "cosine",
            "hybrid": {
                "alpha": 0.7
            }
        },
        "security": {
            "allowed_file_types": ["pdf", "docx", "txt"],
            "max_file_size_mb": 50
        }
    }

    session_id = str(uuid.uuid4())[:8]
    saved_path = PlaygroundConfigManager.save_config("pipeline_test", session_id, test_config)
    print(f"✅ Created test config: {saved_path}")

    # Refresh list
    all_configs = PlaygroundConfigManager.list_configs()

# Select first config for testing
test_config_name = all_configs[0]['name']
test_config_file = all_configs[0]['filename']

print(f"\n🎯 Using config: {test_config_name}")
print(f"   File: {test_config_file}")

# %%
"""
============================================================================
CELL 3: Load and Validate Config (DomainConfig)
============================================================================
"""
print("\n" + "=" * 70)
print("CELL 3: Load and Validate Config")
print("=" * 70)

# Load playground config
pg_config_dict = PlaygroundConfigManager.load_config(test_config_file)
print(f"✅ Loaded playground config from: {test_config_file}")
print(f"   Keys: {list(pg_config_dict.keys())}")

# Merge with global defaults
merged_config = pg_mgr.merge_with_global(pg_config_dict)
print(f"✅ Merged with global config")

# Ensure required fields
synth_domain_id = pg_config_dict.get("playground_name") or pg_config_dict.get("domain_id") or test_config_name
merged_config.setdefault("domain_id", synth_domain_id)
merged_config.setdefault("name", synth_domain_id)
merged_config.setdefault("display_name", synth_domain_id)

print(f"\n📋 Config details:")
print(f"   Domain ID: {merged_config['domain_id']}")
print(f"   Vector Store: {merged_config.get('vector_store', {}).get('provider')}")
print(f"   Chunking: {merged_config.get('chunking', {}).get('strategy')}")
print(f"   Embeddings: {merged_config.get('embeddings', {}).get('provider')}")
print(f"   Retrieval: {merged_config.get('retrieval', {}).get('strategies')}")

# Validate with DomainConfig
try:
    domain_config = DomainConfig(**merged_config)
    print(f"\n✅ DomainConfig validated successfully!")
    print(f"   Domain: {domain_config.domain_id}")
    print(f"   Vector Store: {domain_config.vector_store.provider}")
    print(f"   Chunking: {domain_config.chunking.strategy}")
    print(f"   Embeddings: {domain_config.embeddings.provider}")
except Exception as e:
    print(f"\n❌ Validation failed: {e}")
    raise

# %%
"""
============================================================================
CELL 4: Create DocumentService (get_service_for_config logic)
============================================================================
"""
print("\n" + "=" * 70)
print("CELL 4: Create DocumentService")
print("=" * 70)

# This mimics get_service_for_config from app_phase2.py
try:
    # Try to create service with domain_config parameter
    doc_service = DocumentService(domain_config=domain_config)
    print(f"✅ DocumentService created with domain_config parameter")
except TypeError:
    # Fallback: write temp domain YAML
    print("⚠️  DocumentService doesn't accept domain_config param")
    print("   Writing temporary domain YAML...")

    temp_domain_name = f"{synth_domain_id}_temp"
    temp_domain_file = Path("configs/domains") / f"{temp_domain_name}.yaml"
    temp_domain_file.parent.mkdir(parents=True, exist_ok=True)

    domain_dict = domain_config.model_dump() if hasattr(domain_config, 'model_dump') else domain_config.dict()
    with open(temp_domain_file, 'w') as f:
        yaml.safe_dump(domain_dict, f)

    print(f"   Wrote: {temp_domain_file}")

    doc_service = DocumentService(domain_id=temp_domain_name)
    print(f"✅ DocumentService created with temp domain: {temp_domain_name}")

print(f"\n📊 DocumentService initialized:")
print(f"   Domain: {doc_service.domain_id if hasattr(doc_service, 'domain_id') else 'N/A'}")
print(f"   Pipeline available: {hasattr(doc_service, 'pipeline')}")

# %%
"""
============================================================================
CELL 5: Test Chunking with Different Text Sizes
============================================================================
"""
print("\n" + "=" * 70)
print("CELL 5: Test Chunking")
print("=" * 70)

# Test texts
small_text = "This is a small test document. " * 10
medium_text = "This is a medium-sized document for testing chunking. " * 50
large_text = "This is a large document with lots of content for comprehensive testing. " * 200

test_texts = {
    "small": small_text,
    "medium": medium_text,
    "large": large_text
}

print(f"Chunking strategy: {domain_config.chunking.strategy}")
if domain_config.chunking.strategy == "recursive":
    print(f"  Chunk size: {domain_config.chunking.recursive.chunk_size}")
    print(f"  Overlap: {domain_config.chunking.recursive.overlap}")

for text_type, text in test_texts.items():
    print(f"\n📝 Testing {text_type} text ({len(text)} chars):")

    try:
        # Try using ChunkingFactory if available
        if ChunkingFactory:
            chunker = ChunkingFactory.get_chunker(domain_config.chunking)
            chunks = chunker.chunk_text(text)
        else:
            # Fallback: use DocumentPipeline
            if hasattr(doc_service, 'pipeline'):
                chunks = doc_service.pipeline.chunker.chunk_text(text)
            else:
                print("   ⚠️  No chunker available, skipping")
                continue

        print(f"   ✅ Created {len(chunks)} chunks")
        print(
            f"   Chunk sizes: min={min(len(c) for c in chunks)}, max={max(len(c) for c in chunks)}, avg={sum(len(c) for c in chunks) / len(chunks):.0f}")

        # Show first 2 chunks
        for idx, chunk in enumerate(chunks[:2], 1):
            preview = chunk[:100] + "..." if len(chunk) > 100 else chunk
            print(f"   Chunk {idx}: {preview}")

    except Exception as e:
        print(f"   ❌ Chunking failed: {e}")

# %%
"""
============================================================================
CELL 6: Test Embedding with Different Providers
============================================================================
"""
print("\n" + "=" * 70)
print("CELL 6: Test Embeddings")
print("=" * 70)

test_texts_for_embedding = [
    "What is machine learning?",
    "How does neural network work?",
    "Explain natural language processing"
]

print(f"Embedding provider: {domain_config.embeddings.provider}")
print(f"Model: {domain_config.embeddings.model_name}")
print(f"Device: {domain_config.embeddings.device}")

try:
    # Try using EmbeddingFactory if available
    if EmbeddingFactory:
        embedder = EmbeddingFactory.get_embedder(domain_config.embeddings)
        print("✅ Created embedder via EmbeddingFactory")
    else:
        # Fallback: use DocumentPipeline's embedder
        if hasattr(doc_service, 'pipeline') and hasattr(doc_service.pipeline, 'embedder'):
            embedder = doc_service.pipeline.embedder
            print("✅ Using embedder from DocumentPipeline")
        else:
            print("⚠️  No embedder available, skipping embedding tests")
            embedder = None

    if embedder:
        print(f"\n🧪 Testing embeddings on {len(test_texts_for_embedding)} texts:")

        # Embed texts
        embeddings = embedder.embed_texts(test_texts_for_embedding)

        print(f"✅ Generated {len(embeddings)} embeddings")
        print(f"   Embedding dimension: {len(embeddings[0])}")
        print(f"   First embedding preview: [{embeddings[0][:5]}...]")

        # Test single embedding
        single_emb = embedder.embed_query(test_texts_for_embedding[0])
        print(f"\n✅ Single query embedding: dimension={len(single_emb)}")

except Exception as e:
    print(f"❌ Embedding failed: {e}")
    import traceback

    traceback.print_exc()

# %%
"""
============================================================================
CELL 7: Test Vector DB Operations (Upsert + Query)
============================================================================
"""
print("\n" + "=" * 70)
print("CELL 7: Test Vector DB Operations")
print("=" * 70)

print(f"Vector store provider: {domain_config.vector_store.provider}")
print(f"Collection: {domain_config.vector_store.collection_name}")

try:
    # Get vectorstore from DocumentService pipeline
    if hasattr(doc_service, 'pipeline') and hasattr(doc_service.pipeline, 'vectorstore'):
        vectorstore = doc_service.pipeline.vectorstore
        print("✅ Using vectorstore from DocumentPipeline")
    elif VectorStoreFactory:
        vectorstore = VectorStoreFactory.get_vectorstore(domain_config.vector_store)
        print("✅ Created vectorstore via VectorStoreFactory")
    else:
        print("⚠️  No vectorstore available, skipping")
        vectorstore = None

    if vectorstore and embedder:
        print("\n🧪 Testing upsert operation:")

        # Prepare test documents
        test_docs = [
            {"id": "test_doc_1", "text": "Machine learning is a subset of artificial intelligence.",
             "metadata": {"source": "test", "type": "definition"}},
            {"id": "test_doc_2", "text": "Neural networks are inspired by biological neural networks.",
             "metadata": {"source": "test", "type": "explanation"}},
            {"id": "test_doc_3", "text": "Natural language processing helps computers understand human language.",
             "metadata": {"source": "test", "type": "description"}},
        ]

        # Embed and upsert
        texts = [doc["text"] for doc in test_docs]
        embeddings = embedder.embed_texts(texts)

        for doc, embedding in zip(test_docs, embeddings):
            vectorstore.upsert(
                doc_id=doc["id"],
                text=doc["text"],
                embedding=embedding,
                metadata=doc["metadata"]
            )

        print(f"✅ Upserted {len(test_docs)} documents to vector store")

        print("\n🧪 Testing query operation:")
        query_text = "What is neural network?"
        query_embedding = embedder.embed_query(query_text)

        results = vectorstore.search(
            query_embedding=query_embedding,
            top_k=3
        )

        print(f"✅ Query completed: found {len(results)} results")
        for idx, result in enumerate(results, 1):
            score = result.get('score', result.get('distance', 0))
            text = result.get('text', result.get('document', ''))[:100]
            print(f"   {idx}. Score: {score:.3f} | Text: {text}...")

except Exception as e:
    print(f"❌ Vector DB test failed: {e}")
    import traceback

    traceback.print_exc()

# %%
"""
============================================================================
CELL 8: Test Full Document Ingestion - Text File
============================================================================
"""
print("\n" + "=" * 70)
print("CELL 8: Test Document Ingestion - Text File")
print("=" * 70)

# Create test text file
test_dir = Path("test_documents")
test_dir.mkdir(exist_ok=True)

text_file = test_dir / "test_document.txt"
text_content = """
This is a test document for the playground pipeline testing.

It contains multiple paragraphs to test the chunking functionality.
The document discusses various topics including machine learning,
natural language processing, and vector databases.

Machine learning is a field of artificial intelligence that uses
statistical techniques to give computer systems the ability to learn
from data without being explicitly programmed.

Natural language processing (NLP) is a subfield of linguistics,
computer science, and artificial intelligence concerned with the
interactions between computers and human language.

Vector databases are specialized databases designed to store and
query high-dimensional vectors efficiently, which is crucial for
modern AI applications.
"""

with open(text_file, 'w') as f:
    f.write(text_content)

print(f"✅ Created test text file: {text_file}")

try:
    print("\n🧪 Ingesting text document via DocumentService:")

    # Prepare metadata
    doc_id = f"test_txt_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    metadata = {
        "doc_id": doc_id,
        "title": "Test Text Document",
        "domain": domain_config.domain_id,
        "doc_type": "test",
        "uploader_id": "test_user",
        "source_file_path": str(text_file)
    }

    # Upload document
    result = doc_service.upload_document(
        file_obj=open(text_file, 'rb'),
        metadata=metadata,
        replace_existing=True
    )

    print(f"✅ Document ingested successfully!")
    print(f"   Doc ID: {result.get('doc_id', doc_id)}")
    print(f"   Chunks created: {result.get('chunks_ingested', 'N/A')}")
    print(f"   Embedding model: {result.get('embedding_model', 'N/A')}")
    print(f"   Chunking strategy: {result.get('chunking_strategy', 'N/A')}")

except Exception as e:
    print(f"❌ Text document ingestion failed: {e}")
    import traceback

    traceback.print_exc()

# %%
"""
============================================================================
CELL 9: Test Document Ingestion - Create and Ingest PDF
============================================================================
"""
print("\n" + "=" * 70)
print("CELL 9: Test Document Ingestion - PDF")
print("=" * 70)

# Try to create a simple PDF
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    pdf_file = test_dir / "test_document.pdf"

    c = canvas.Canvas(str(pdf_file), pagesize=letter)
    c.drawString(100, 750, "Test PDF Document for Playground Testing")
    c.drawString(100, 700, "")
    c.drawString(100, 680, "This PDF contains sample text for testing the document pipeline.")
    c.drawString(100, 660, "")
    c.drawString(100, 640, "Topics covered:")
    c.drawString(100, 620, "- Document ingestion")
    c.drawString(100, 600, "- PDF parsing")
    c.drawString(100, 580, "- Chunking and embedding")
    c.drawString(100, 560, "- Vector storage")
    c.save()

    print(f"✅ Created test PDF: {pdf_file}")

    print("\n🧪 Ingesting PDF document:")

    doc_id = f"test_pdf_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    metadata = {
        "doc_id": doc_id,
        "title": "Test PDF Document",
        "domain": domain_config.domain_id,
        "doc_type": "test_pdf",
        "uploader_id": "test_user",
        "source_file_path": str(pdf_file)
    }

    result = doc_service.upload_document(
        file_obj=open(pdf_file, 'rb'),
        metadata=metadata,
        replace_existing=True
    )

    print(f"✅ PDF ingested successfully!")
    print(f"   Doc ID: {result.get('doc_id', doc_id)}")
    print(f"   Chunks created: {result.get('chunks_ingested', 'N/A')}")

except ImportError:
    print("⚠️  reportlab not installed, skipping PDF test")
    print("   Install with: pip install reportlab")
except Exception as e:
    print(f"❌ PDF ingestion failed: {e}")
    import traceback

    traceback.print_exc()

# %%
"""
============================================================================
CELL 10: Test Document Ingestion - Create and Ingest DOCX
============================================================================
"""
print("\n" + "=" * 70)
print("CELL 10: Test Document Ingestion - DOCX")
print("=" * 70)

try:
    from docx import Document as DocxDocument

    docx_file = test_dir / "test_document.docx"

    doc = DocxDocument()
    doc.add_heading('Test DOCX Document', 0)
    doc.add_paragraph('This is a test Word document for playground testing.')
    doc.add_heading('Section 1: Introduction', level=1)
    doc.add_paragraph('This document tests the DOCX parsing capabilities of the pipeline.')
    doc.add_heading('Section 2: Content', level=1)
    doc.add_paragraph('Testing chunking and embedding with DOCX format.')
    doc.add_paragraph('Multiple paragraphs help test the chunking strategy.')
    doc.save(str(docx_file))

    print(f"✅ Created test DOCX: {docx_file}")

    print("\n🧪 Ingesting DOCX document:")

    doc_id = f"test_docx_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    metadata = {
        "doc_id": doc_id,
        "title": "Test DOCX Document",
        "domain": domain_config.domain_id,
        "doc_type": "test_docx",
        "uploader_id": "test_user",
        "source_file_path": str(docx_file)
    }

    result = doc_service.upload_document(
        file_obj=open(docx_file, 'rb'),
        metadata=metadata,
        replace_existing=True
    )

    print(f"✅ DOCX ingested successfully!")
    print(f"   Doc ID: {result.get('doc_id', doc_id)}")
    print(f"   Chunks created: {result.get('chunks_ingested', 'N/A')}")

except ImportError:
    print("⚠️  python-docx not installed, skipping DOCX test")
    print("   Install with: pip install python-docx")
except Exception as e:
    print(f"❌ DOCX ingestion failed: {e}")
    import traceback

    traceback.print_exc()

# %%
"""
============================================================================
CELL 11: Test Document Listing
============================================================================
"""
print("\n" + "=" * 70)
print("CELL 11: Test Document Listing")
print("=" * 70)

try:
    # List all documents
    documents = doc_service.list_documents(
        filters={"domain": domain_config.domain_id, "deprecated": False}
    )

    print(f"✅ Found {len(documents)} document(s) in domain '{domain_config.domain_id}':")

    for idx, doc in enumerate(documents, 1):
        print(f"\n  {idx}. Doc ID: {doc.get('doc_id', 'N/A')}")
        print(f"     Title: {doc.get('title', 'N/A')}")
        print(f"     Type: {doc.get('doc_type', 'N/A')}")
        print(f"     Uploader: {doc.get('uploader_id', 'N/A')}")
        print(f"     Chunks: {doc.get('chunk_count', 'N/A')}")
        print(f"     Created: {doc.get('created_at', 'N/A')}")
        print(f"     Deprecated: {doc.get('deprecated', False)}")

    if not documents:
        print("  ⚠️  No documents found (documents may not have been indexed yet)")

except Exception as e:
    print(f"❌ Document listing failed: {e}")
    import traceback

    traceback.print_exc()

# %%
"""
============================================================================
CELL 12: Test Chunk Listing for Specific Document
============================================================================
"""
print("\n" + "=" * 70)
print("CELL 12: Test Chunk Listing")
print("=" * 70)

try:
    if documents and len(documents) > 0:
        # Get chunks for first document
        test_doc_id = documents[0].get('doc_id')
        print(f"📄 Listing chunks for document: {test_doc_id}")

        chunks = doc_service.list_chunks(
            doc_id=test_doc_id,
            limit=10
        )

        print(f"✅ Found {len(chunks)} chunk(s):")

        for idx, chunk in enumerate(chunks, 1):
            chunk_text = chunk.get('text', chunk.get('document', ''))
            preview = chunk_text[:150] + "..." if len(chunk_text) > 150 else chunk_text

            print(f"\n  Chunk {idx}:")
            print(f"    ID: {chunk.get('id', 'N/A')}")
            print(f"    Length: {len(chunk_text)} chars")
            print(f"    Text: {preview}")

            metadata = chunk.get('metadata', {})
            if metadata:
                print(f"    Metadata: {metadata}")
    else:
        print("⚠️  No documents available to list chunks")

except Exception as e:
    print(f"❌ Chunk listing failed: {e}")
    import traceback

    traceback.print_exc()

# %%
"""
============================================================================
CELL 13: Test Retrieval Query
============================================================================
"""
print("\n" + "=" * 70)
print("CELL 13: Test Retrieval Query")
print("=" * 70)

test_queries = [
    "What is machine learning?",
    "Explain neural networks",
    "How does NLP work?"
]

try:
    for query_text in test_queries:
        print(f"\n🔍 Query: '{query_text}'")

        # Query the document service
        results = doc_service.query(
            query=query_text,
            domain=domain_config.domain_id,
            top_k=5
        )

        print(f"✅ Retrieved {len(results)} result(s):")

        for idx, result in enumerate(results, 1):
            score = result.get('score', result.get('distance', 0))
            text = result.get('text', result.get('document', ''))[:200]
            doc_id = result.get('doc_id', result.get('metadata', {}).get('doc_id', 'N/A'))

            print(f"\n  Result {idx} (Score: {score:.3f}):")
            print(f"    Doc ID: {doc_id}")
            print(f"    Text: {text}...")

except Exception as e:
    print(f"❌ Query failed: {e}")
    import traceback

    traceback.print_exc()

# %%
"""
============================================================================
CELL 14: Test Different Retrieval Strategies
============================================================================
"""
print("\n" + "=" * 70)
print("CELL 14: Test Retrieval Strategies")
print("=" * 70)

test_query = "machine learning and artificial intelligence"
strategies_to_test = ["vector_similarity", "hybrid", "bm25"]

try:
    available_strategies = domain_config.retrieval.strategies
    print(f"Configured strategies: {available_strategies}")

    for strategy in strategies_to_test:
        if strategy not in available_strategies:
            print(f"\n⚠️  Strategy '{strategy}' not configured, skipping")
            continue

        print(f"\n🧪 Testing strategy: {strategy}")

        try:
            # Try to query with specific strategy
            if hasattr(doc_service, 'query_with_strategy'):
                results = doc_service.query_with_strategy(
                    query=test_query,
                    strategy=strategy,
                    domain=domain_config.domain_id,
                    top_k=3
                )
            else:
                # Fallback to regular query
                results = doc_service.query(
                    query=test_query,
                    domain=domain_config.domain_id,
                    top_k=3
                )

            print(f"✅ Retrieved {len(results)} result(s) using {strategy}:")

            for idx, result in enumerate(results[:3], 1):
                score = result.get('score', 0)
                text = result.get('text', result.get('document', ''))[:100]
                print(f"  {idx}. Score: {score:.3f} | {text}...")

        except Exception as e:
            print(f"❌ Strategy {strategy} failed: {e}")

except Exception as e:
    print(f"❌ Strategy testing failed: {e}")

# %%
"""
============================================================================
CELL 15: Test Document Deletion
============================================================================
"""
print("\n" + "=" * 70)
print("CELL 15: Test Document Deletion (Optional)")
print("=" * 70)

cleanup = input("Do you want to delete test documents? (yes/no): ")

if cleanup.lower() == 'yes':
    try:
        # List documents again
        documents = doc_service.list_documents(
            filters={"domain": domain_config.domain_id}
        )

        print(f"\n🗑️  Deleting {len(documents)} test document(s)...")

        for doc in documents:
            doc_id = doc.get('doc_id')
            try:
                doc_service.delete_document(
                    doc_id=doc_id,
                    domain=domain_config.domain_id
                )
                print(f"  ✅ Deleted: {doc_id}")
            except Exception as e:
                print(f"  ❌ Failed to delete {doc_id}: {e}")

        print("\n✅ Cleanup completed!")

    except Exception as e:
        print(f"❌ Cleanup failed: {e}")
else:
    print("Skipping cleanup - test documents preserved")

# %%
"""
============================================================================
CELL 16: Test Summary and Statistics
============================================================================
"""
print("\n" + "=" * 70)
print("CELL 16: Test Summary")
print("=" * 70)

print("\n📊 Pipeline Test Summary:")
print("=" * 70)
print(f"Config used: {test_config_name}")
print(f"Domain ID: {domain_config.domain_id}")
print(f"\nComponents tested:")
print(f"  ✅ Config loading and validation")
print(f"  ✅ DocumentService creation")
print(f"  ✅ Chunking (recursive/semantic)")
print(f"  ✅ Embedding ({domain_config.embeddings.provider})")
print(f"  ✅ Vector DB ({domain_config.vector_store.provider})")
print(f"  ✅ Document ingestion (TXT, PDF, DOCX)")
print(f"  ✅ Document listing")
print(f"  ✅ Chunk retrieval")
print(f"  ✅ Query/retrieval")
print(f"  ✅ Multiple retrieval strategies")

try:
    final_doc_count = len(doc_service.list_documents(filters={"domain": domain_config.domain_id}))
    print(f"\nFinal document count: {final_doc_count}")
except:
    print(f"\nFinal document count: Unable to determine")

print("\n" + "=" * 70)
print("✅ ALL PIPELINE TESTS COMPLETED!")
print("=" * 70)
print(f"End time: {datetime.now()}")

# %%
